import os, json, uuid, requests, io, zipfile
from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS
from groq import Groq
from PIL import Image, ImageDraw, ImageFont, ImageColor
import arabic_reshaper
from bidi.algorithm import get_display
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)
CORS(app)

GROQ_API_KEY = os.environ.get('GROQ_API_KEY', '')
groq_client = Groq(api_key=GROQ_API_KEY) if GROQ_API_KEY else None

# ========== SUPPORTED LANGUAGES ==========
LANGUAGES = {
    'arabic': {'name': 'العربية', 'code': 'ar', 'direction': 'rtl', 'font': 'Arial'},
    'english': {'name': 'English', 'code': 'en', 'direction': 'ltr', 'font': 'Helvetica'},
    'french': {'name': 'Français', 'code': 'fr', 'direction': 'ltr', 'font': 'Helvetica'},
    'spanish': {'name': 'Español', 'code': 'es', 'direction': 'ltr', 'font': 'Helvetica'},
    'german': {'name': 'Deutsch', 'code': 'de', 'direction': 'ltr', 'font': 'Helvetica'},
    'italian': {'name': 'Italiano', 'code': 'it', 'direction': 'ltr', 'font': 'Helvetica'},
    'portuguese': {'name': 'Português', 'code': 'pt', 'direction': 'ltr', 'font': 'Helvetica'},
    'russian': {'name': 'Русский', 'code': 'ru', 'direction': 'ltr', 'font': 'Helvetica'},
    'chinese': {'name': '中文', 'code': 'zh', 'direction': 'ltr', 'font': 'Helvetica'},
    'japanese': {'name': '日本語', 'code': 'ja', 'direction': 'ltr', 'font': 'Helvetica'},
    'korean': {'name': '한국어', 'code': 'ko', 'direction': 'ltr', 'font': 'Helvetica'},
    'hindi': {'name': 'हिन्दी', 'code': 'hi', 'direction': 'ltr', 'font': 'Helvetica'},
    'turkish': {'name': 'Türkçe', 'code': 'tr', 'direction': 'ltr', 'font': 'Helvetica'},
    'urdu': {'name': 'اردو', 'code': 'ur', 'direction': 'rtl', 'font': 'Arial'},
    'malay': {'name': 'Bahasa Melayu', 'code': 'ms', 'direction': 'ltr', 'font': 'Helvetica'},
}

# ========== COVER THEMES ==========
COVER_THEMES = {
    'educational': {'bg': '#1a237e', 'accent': '#ffd600', 'icon': '📚'},
    'story': {'bg': '#4a148c', 'accent': '#ff4081', 'icon': '📖'},
    'self_help': {'bg': '#004d40', 'accent': '#00e5ff', 'icon': '💪'},
    'tech': {'bg': '#0d47a1', 'accent': '#00e676', 'icon': '💻'},
    'religious': {'bg': '#1b5e20', 'accent': '#ffd700', 'icon': '🕌'},
    'history': {'bg': '#3e2723', 'accent': '#ffab00', 'icon': '📜'},
    'cooking': {'bg': '#bf360c', 'accent': '#ffeb3b', 'icon': '🍳'},
    'business': {'bg': '#01579b', 'accent': '#ff6d00', 'icon': '💼'},
    'science': {'bg': '#311b92', 'accent': '#00e5ff', 'icon': '🔬'},
    'romance': {'bg': '#880e4f', 'accent': '#ff80ab', 'icon': '💕'},
}

def process_text(text, lang_code):
    """Process text based on language direction"""
    if LANGUAGES.get(lang_code, {}).get('direction') == 'rtl':
        reshaped = arabic_reshaper.reshape(text)
        return get_display(reshaped)
    return text

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/generate-book', methods=['POST'])
def generate_book():
    try:
        data = request.get_json(force=True)
        topic = data.get('topic', '')
        genre = data.get('genre', 'educational')
        chapters = min(int(data.get('chapters', 5)), 10)
        language = data.get('language', 'arabic')
        author = data.get('author', 'Author')
        max_pages = min(int(data.get('max_pages', 100)), 100)
        
        if not topic or not groq_client:
            return jsonify({"error": "Topic and API key required"}), 400
        
        lang_info = LANGUAGES.get(language, LANGUAGES['english'])
        
        # Generate book structure
        prompt = f"""Write a complete {genre} book in {lang_info['name']} about "{topic}".
        Author: {author}
        Maximum pages: {max_pages}
        
        CRITICAL RULES:
        1. Each chapter must have COMPLETELY DIFFERENT content
        2. NO repetition between chapters
        3. NO spelling or grammar errors
        4. Write REAL, accurate content about the topic
        5. Each chapter should be 300-500 words
        
        Return ONLY valid JSON:
        {{
            "title": "Book Title",
            "subtitle": "Subtitle",
            "description": "2 sentence description",
            "chapters": [
                {{"number":1,"title":"Chapter Title","content":"Full chapter content..."}}
            ]
        }}
        """
        
        response = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role":"user","content":prompt}],
            temperature=0.8,
            max_tokens=8000
        )
        
        text = response.choices[0].message.content
        start = text.find('{')
        end = text.rfind('}') + 1
        book_data = json.loads(text[start:end])
        
        # Ensure we have chapters
        if not book_data.get('chapters'):
            book_data['chapters'] = [{"number":1,"title":"Introduction","content":"Content here..."}]
        
        # Limit chapters
        book_data['chapters'] = book_data['chapters'][:chapters]
        
        # Generate cover
        cover_path = create_cover(book_data['title'], book_data.get('subtitle',''), author, genre, language)
        
        # Create downloads
        book_id = uuid.uuid4().hex[:12]
        os.makedirs(f'static/books/{book_id}', exist_ok=True)
        os.makedirs('static/covers', exist_ok=True)
        
        # PDF
        pdf_path = create_pdf(book_data, author, book_id, language)
        # DOCX  
        docx_path = create_docx(book_data, author, book_id)
        # TXT
        txt_path = create_txt(book_data, author, book_id)
        # HTML
        html_path = create_html(book_data, author, book_id)
        
        # ZIP
        zip_path = f'static/books/{book_id}/book_formats.zip'
        with zipfile.ZipFile(zip_path, 'w') as zf:
            zf.write(pdf_path, 'book.pdf')
            zf.write(docx_path, 'book.docx')
            zf.write(txt_path, 'book.txt')
            zf.write(html_path, 'book.html')
        
        return jsonify({
            "success": True,
            "book_id": book_id,
            "title": book_data['title'],
            "subtitle": book_data.get('subtitle', ''),
            "description": book_data.get('description', ''),
            "chapters": book_data['chapters'],
            "cover_url": f"/{cover_path}",
            "downloads": {
                "pdf": f"/{pdf_path}",
                "docx": f"/{docx_path}",
                "txt": f"/{txt_path}",
                "html": f"/{html_path}",
                "zip": f"/{zip_path}"
            }
        })
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

def create_cover(title, subtitle, author, genre, language):
    """Create professional book cover"""
    theme = COVER_THEMES.get(genre, COVER_THEMES['educational'])
    
    # Create image
    img = Image.new('RGB', (800, 1200), theme['bg'])
    draw = ImageDraw.Draw(img)
    
    # Decorative pattern
    for i in range(0, 800, 40):
        draw.line([(i, 0), (i+20, 1200)], fill=theme['bg'], width=2)
    
    # Gradient overlay
    for y in range(600):
        alpha = int(255 * (1 - y/600))
        color = ImageColor.getrgb(theme['accent'])
        draw.line([(0, y), (800, y)], fill=(*color, alpha))
    
    # Border
    draw.rectangle([20, 20, 780, 1180], outline=theme['accent'], width=3)
    draw.rectangle([30, 30, 770, 1170], outline=theme['accent'], width=1)
    
    # Title (centered)
    try:
        font_large = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 48)
        font_medium = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 28)
        font_small = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 20)
    except:
        font_large = ImageFont.load_default()
        font_medium = ImageFont.load_default()
        font_small = ImageFont.load_default()
    
    # Process text for RTL languages
    display_title = process_text(title, LANGUAGES.get(language, {}).get('code', 'en'))
    
    # Draw title
    y_pos = 200
    words = display_title.split()
    for word in words:
        bbox = draw.textbbox((0, 0), word, font=font_large)
        text_width = bbox[2] - bbox[0]
        x_pos = (800 - text_width) / 2
        draw.text((x_pos, y_pos), word, fill=theme['accent'], font=font_large)
        y_pos += 55
    
    # Subtitle
    if subtitle:
        display_sub = process_text(subtitle, LANGUAGES.get(language, {}).get('code', 'en'))
        y_pos += 30
        bbox = draw.textbbox((0, 0), display_sub, font=font_medium)
        text_width = bbox[2] - bbox[0]
        draw.text(((800-text_width)/2, y_pos), display_sub, fill='#ffffff', font=font_medium)
    
    # Author
    display_author = process_text(f'By: {author}', 'en')
    bbox = draw.textbbox((0, 0), display_author, font=font_small)
    text_width = bbox[2] - bbox[0]
    draw.text(((800-text_width)/2, 1050), display_author, fill='#ffffff', font=font_small)
    
    # Genre icon
    bbox = draw.textbbox((0, 0), theme['icon'], font=font_large)
    text_width = bbox[2] - bbox[0]
    draw.text(((800-text_width)/2, 1100), theme['icon'], fill=theme['accent'], font=font_large)
    
    path = f'static/covers/cover_{uuid.uuid4().hex[:8]}.png'
    img.save(path)
    return path

def create_pdf(book_data, author, book_id, language):
    """Create PDF with proper formatting"""
    path = f'static/books/{book_id}/book.pdf'
    c = canvas.Canvas(path, pagesize=A4)
    width, height = A4
    
    # Cover page
    c.setFont("Helvetica-Bold", 28)
    c.drawCentredString(width/2, height-200, book_data['title'])
    if book_data.get('subtitle'):
        c.setFont("Helvetica", 16)
        c.drawCentredString(width/2, height-240, book_data['subtitle'])
    c.setFont("Helvetica", 12)
    c.drawCentredString(width/2, height-280, f"By: {author}")
    c.showPage()
    
    # Table of Contents
    c.setFont("Helvetica-Bold", 20)
    c.drawString(50, height-100, "Table of Contents")
    c.line(50, height-110, width-50, height-110)
    
    y = height - 140
    for ch in book_data['chapters']:
        c.setFont("Helvetica", 12)
        c.drawString(70, y, f"Chapter {ch['number']}: {ch['title']}")
        y -= 25
    
    # Chapters
    page_count = 2
    for ch in book_data['chapters']:
        if page_count >= 100:
            break
            
        c.showPage()
        page_count += 1
        
        c.setFont("Helvetica-Bold", 18)
        c.drawString(50, height-100, f"Chapter {ch['number']}")
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, height-130, ch['title'])
        c.line(50, height-140, width-50, height-140)
        
        # Content
        c.setFont("Helvetica", 11)
        y = height - 170
        content = ch.get('content', '')
        words = content.split()
        line = ""
        
        for word in words:
            test_line = line + " " + word if line else word
            if c.stringWidth(test_line, "Helvetica", 11) < width - 100:
                line = test_line
            else:
                if y < 50:
                    if page_count >= 100:
                        break
                    c.showPage()
                    page_count += 1
                    c.setFont("Helvetica", 11)
                    y = height - 50
                c.drawString(50, y, line)
                y -= 18
                line = word
        
        if line and y >= 50 and page_count < 100:
            c.drawString(50, y, line)
    
    c.save()
    return path

def create_docx(book_data, author, book_id):
    """Create Word document"""
    path = f'static/books/{book_id}/book.docx'
    doc = Document()
    
    # Title page
    doc.add_heading(book_data['title'], 0)
    if book_data.get('subtitle'):
        doc.add_heading(book_data['subtitle'], 1)
    doc.add_paragraph(f'By: {author}')
    doc.add_page_break()
    
    # TOC
    doc.add_heading('Table of Contents', 1)
    for ch in book_data['chapters']:
        doc.add_paragraph(f"Chapter {ch['number']}: {ch['title']}")
    
    # Chapters
    for ch in book_data['chapters']:
        doc.add_page_break()
        doc.add_heading(f"Chapter {ch['number']}: {ch['title']}", 1)
        doc.add_paragraph(ch.get('content', ''))
    
    doc.save(path)
    return path

def create_txt(book_data, author, book_id):
    """Create text file"""
    path = f'static/books/{book_id}/book.txt'
    with open(path, 'w', encoding='utf-8') as f:
        f.write(f"{book_data['title']}\\n")
        f.write(f"By: {author}\\n")
        f.write("="*50 + "\\n\\n")
        for ch in book_data['chapters']:
            f.write(f"CHAPTER {ch['number']}: {ch['title']}\\n")
            f.write("-"*30 + "\\n")
            f.write(ch.get('content', '') + "\\n\\n")
    return path

def create_html(book_data, author, book_id):
    """Create HTML file"""
    path = f'static/books/{book_id}/book.html'
    with open(path, 'w', encoding='utf-8') as f:
        f.write(f"""<!DOCTYPE html>
<html>
<head><meta charset="UTF-8"><title>{book_data['title']}</title>
<style>
body{{font-family:Arial,sans-serif;max-width:800px;margin:0 auto;padding:20px;line-height:1.8}}
h1{{color:#1a237e;text-align:center}}
h2{{color:#4a148c;margin-top:40px}}
.chapter{{margin:30px 0;padding:20px;background:#f5f5f5;border-radius:10px}}
</style></head>
<body>
<h1>{book_data['title']}</h1>
<p style="text-align:center">By: {author}</p>
<hr>
""")
        for ch in book_data['chapters']:
            f.write(f'<div class="chapter"><h2>Chapter {ch["number"]}: {ch["title"]}</h2><p>{ch.get("content","")}</p></div>')
        f.write('</body></html>')
    return path

@app.route('/api/health')
def health():
    return jsonify({
        "status": "ok",
        "groq": bool(GROQ_API_KEY),
        "languages": list(LANGUAGES.keys())
    })

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
